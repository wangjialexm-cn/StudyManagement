from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "/Users/user/Documents/家庭教师/五天英语补缺练习册（学生版）.docx"
FONT = "PingFang SC"
BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "E8EEF5"
PALE = "F4F7FB"
GRAY = "667085"
BLACK = "111111"


def set_font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=GRAY)
    fld = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    fld._r.extend([begin, instr, end])
    set_font(fld, 9, color=GRAY)
    tail = paragraph.add_run(" 页")
    set_font(tail, 9, color=GRAY)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        st = doc.styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for style_name in ("List Number", "List Bullet"):
        st = doc.styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.19)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.line_spacing = 1.18


def set_header_footer(section):
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("五天英语补缺练习册  |  轻量恢复期")
    set_font(r, 9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, 26, True, DARK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(26)
        set_font(p2.add_run(subtitle), 14, color=GRAY)


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + "  "), 10.5, True, DARK)
    set_font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def info_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("姓名：__________________    日期：__________________    用时：________ 分钟"), 10.5)


def heading(doc, text, level=1):
    doc.add_paragraph(text, style=f"Heading {level}")


def body(doc, text, bold=False, color=BLACK, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    set_font(p.add_run(text), 11, bold, color)
    return p


def numbered(doc, text, lines=0):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(text), 11)
    for _ in range(lines):
        q = doc.add_paragraph("____________________________________________________________________________")
        q.paragraph_format.left_indent = Inches(0.38)
        q.paragraph_format.space_after = Pt(5)
        set_font(q.runs[0], 9, color=GRAY)


def vocab_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    repeat_header(hdr)
    for i, text in enumerate(("题号", "中文提示", "写出英文")):
        shade(hdr.cells[i], LIGHT)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), 10, True, DARK)
    for i, cue in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cells[0].paragraphs[0].add_run(str(i)), 10)
        set_font(cells[1].paragraphs[0].add_run(cue), 10.5)
        set_font(cells[2].paragraphs[0].add_run("________________________________"), 10, color=GRAY)
        if i % 2 == 0:
            for c in cells:
                shade(c, "FAFBFC")
    set_table_geometry(table, [700, 3260, 5400])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def options(doc, question, choices):
    numbered(doc, question)
    p = doc.add_paragraph("    " + "      ".join(choices))
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.runs[0], 10.5)


def day_start(doc, n, title_text, focus):
    if n > 1:
        doc.add_page_break()
    heading(doc, f"第 {n} 天  {title_text}", 1)
    info_line(doc)
    callout(doc, "今日重点", focus + "。建议用时 60 分钟，到点即停。")


def build():
    doc = Document()
    setup_styles(doc)
    set_header_footer(doc.sections[0])

    title(doc, "五天英语补缺练习册", "中考后轻量恢复期 · 每天约 1 小时 · 学生版")
    body(doc, "这份练习根据之前的真实错题和单词复习记录设计。目的不是增加负担，而是用五天把最容易反复失分的基础重新接稳。", after=14)
    callout(doc, "使用方法", "每天完成一份；不会的题可以空着或写“不会”；不要查答案。做完后拍照发给我，我会逐题批改并更新后续安排。")
    heading(doc, "五天安排", 2)
    schedule = [
        ("1", "高频错词 + 过去时态"),
        ("2", "易混词 + 被动语态"),
        ("3", "情态动词 + 情景交际"),
        ("4", "非谓语 + 词形变化 + 填词"),
        ("5", "综合检测 + 短文写作"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(("天数", "内容")):
        shade(table.rows[0].cells[i], LIGHT)
        set_font(table.rows[0].cells[i].paragraphs[0].add_run(h), 10.5, True, DARK)
    for d, content in schedule:
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run("第 " + d + " 天"), 10.5, True)
        set_font(cells[1].paragraphs[0].add_run(content), 10.5)
    set_table_geometry(table, [1800, 7560])
    heading(doc, "完成约定", 2)
    for text in ("独立完成，不查词典和答案。", "遇到不会的题先做记号，继续往下做。", "每天最多 60 分钟，不加量，不熬夜。", "拍照时保持页面完整、文字清晰。"): 
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(text), 11)

    day_start(doc, 1, "高频错词与过去时态", "主动拼写；一般过去时与过去进行时")
    heading(doc, "一、词汇主动拼写（约 20 分钟）", 2)
    vocab_table(doc, ["减少；降低", "保持；仍然", "目的；用途", "图像；形象", "意识到", "研究", "摧毁；破坏", "打败", "激动；兴奋", "可用的；有空的", "迅速的", "准备；准备工作"])
    heading(doc, "二、用所给词的正确形式填空（约 15 分钟）", 2)
    for q in [
        "I ________ (wash) my clothes yesterday.",
        "Lily ________ (sing) at nine yesterday evening.",
        "At that moment, Tom ________ (watch) a documentary.",
        "Long ago, people ________ (climb) the mountain to celebrate the festival.",
        "He ________ (fly) a kite with his father last Sunday.",
        "While she was cooking, her brother ________ (play) the violin.",
        "What ________ you ________ (do) last weekend?",
        "The boy ________ (realize) his mistake after the teacher talked to him.",
    ]:
        numbered(doc, q)
    heading(doc, "三、改错（约 10 分钟）", 2)
    for q in [
        "I wash my clothes yesterday.",
        "I was make a snowman at this time yesterday.",
        "She were reading when I called her.",
        "What did you did last weekend?",
    ]:
        numbered(doc, q, 1)
    heading(doc, "四、看中文写英文（约 15 分钟）", 2)
    for q in ["昨天我放了风筝。", "昨晚九点莉莉正在唱歌。", "那时他正在看一部纪录片。", "我意识到了自己的错误。", "这项研究带来了快速进步。"]:
        numbered(doc, q, 1)

    day_start(doc, 2, "易混词与被动语态", "区分易混词；掌握 be + 过去分词")
    heading(doc, "一、易混词选择（约 15 分钟）", 2)
    pairs = [
        ("I ________ a present from Jack, but I didn't ________ it.", "received / accept"),
        ("Please use the machine ________. It will ________ rain tonight.", "properly / probably"),
        ("Living ________ are improving, and students have more ________ now.", "conditions / confidence"),
        ("This factory ________ many new ________ every year.", "produces / products"),
        ("Please ________ the old chair with a new one, and ________ the broken parts.", "replace / remove"),
        ("Everyone came ________ Tom, because he refused to ________ the invitation.", "except / accept"),
        ("The teacher asked us to ________ the video and not to ________ the same mistake.", "record / repeat"),
        ("The ________ of the activity is to improve our sense of ________.", "purpose / satisfaction"),
    ]
    for sentence, words in pairs:
        numbered(doc, sentence + f"  （选词：{words}）")
    heading(doc, "二、被动语态填空（约 20 分钟）", 2)
    for q in [
        "The documentary ________ (show) in Italy last year.",
        "Nanyin ________ usually ________ (play) with traditional instruments.",
        "The Temple of Heaven ________ (add) to the World Heritage List in 1998.",
        "The old bridge ________ (wash) away by the flood last summer.",
        "The new airport ________ (complete) by the end of this year.",
        "These books ________ (use) by many students every day.",
        "The classroom ________ (clean) yesterday afternoon.",
        "English ________ (speak) in many countries.",
    ]:
        numbered(doc, q)
    heading(doc, "三、单项选择（约 10 分钟）", 2)
    options(doc, "The documentary was ________ in Italy.", ["A. shown", "B. wrote", "C. grew"])
    options(doc, "The work will ________ before Friday.", ["A. complete", "B. be completed", "C. completed"])
    options(doc, "Tea is ________ in many parts of China.", ["A. grown", "B. grew", "C. growing"])
    options(doc, "The electricity supply was ________ because of the storm.", ["A. cut down", "B. cut off", "C. cut up"])
    heading(doc, "四、翻译（约 15 分钟）", 2)
    for q in ["这部纪录片去年在中国放映。", "教室每天都被学生打扫。", "这座桥将在今年年底建成。", "请正确地使用这台机器。", "我收到邀请了，但没有接受它。"]:
        numbered(doc, q, 1)

    day_start(doc, 3, "情态动词与情景交际", "mustn't、needn't、can't；标准交际表达")
    heading(doc, "一、选择正确的情态动词（约 15 分钟）", 2)
    for q, choices in [
        ("Students ________ bring mobile phones into the exam room.", ["A. mustn't", "B. needn't", "C. can"]),
        ("You ________ bring a gift. Just come and have fun.", ["A. mustn't", "B. needn't", "C. can't"]),
        ("The lights are off. Ella ________ be in the classroom.", ["A. must", "B. should", "C. can't"]),
        ("—Must I hand it in now? —No, you ________.", ["A. mustn't", "B. needn't", "C. can't"]),
        ("Grandpa asked Parker if he ________ take a walk.", ["A. can", "B. could", "C. must"]),
        ("You look tired. You ________ have a rest.", ["A. should", "B. can't", "C. mustn't"]),
    ]:
        options(doc, q, choices)
    heading(doc, "二、根据情景写出完整英语（约 25 分钟）", 2)
    for q in [
        "别人邀请你看电影，你愿意参加。",
        "你想请别人打开窗户，并询问对方是否介意。",
        "朋友考试没考好，你表示安慰。",
        "别人向你道歉，你回答“没关系”。",
        "电话中介绍自己是李华。",
        "询问今天的日期。",
        "询问今天星期几。",
        "祝贺朋友在演讲比赛中获一等奖。",
        "感谢朋友邀请你。",
        "询问对方在这座城市住了多久。",
    ]:
        numbered(doc, q, 1)
    heading(doc, "三、改正不自然或错误的回答（约 10 分钟）", 2)
    for q in [
        "—Would you mind opening the window? —No mind.",
        "—I'm sorry I'm late. —Not at all.",
        "—I won first prize! —I'm sorry to hear that.",
        "—Would you like to join us? —Yes, I'd love.",
    ]:
        numbered(doc, q, 1)
    heading(doc, "四、情景小对话（约 10 分钟）", 2)
    body(doc, "补全对话，每空写一句完整的话。")
    body(doc, "A: Hi, Lisa. ____________________________________________?")
    body(doc, "B: I've lived in this city for two years.")
    body(doc, "A: Would you like to go to the movies with me this evening?")
    body(doc, "B: ____________________________________________. What time shall we meet?")
    body(doc, "A: At seven. Please don't be late.")
    body(doc, "B: ____________________________________________.")

    day_start(doc, 4, "非谓语、词形变化与短文填词", "doing 结构；形容词与副词；代词形式")
    heading(doc, "一、用所给词的正确形式填空（约 20 分钟）", 2)
    for q in [
        "Keep ________ (try), and you will make progress.",
        "Kate enjoys ________ (shop) with her mother.",
        "Would you mind ________ (open) the window?",
        "I feel like ________ (take) a walk after dinner.",
        "She is interested in ________ (learn) Chinese culture.",
        "They spent a lot of money ________ (protect) the old building.",
        "He thought of ________ (build) a bridge over the river.",
        "We can learn English by ________ (read) aloud.",
        "Thank you for ________ (invite) me.",
        "The local people ________ (traditional) drink tea at the festival.",
        "She expressed her ideas ________ (confident).",
        "Now I understand the story more ________ (deep).",
        "The characters are simple but very ________ (use).",
        "He completed the work ________ (careful).",
    ]:
        numbered(doc, q)
    heading(doc, "二、选择正确的代词（约 10 分钟）", 2)
    for q, choices in [
        ("The temple is famous for ________ long history.", ["A. it", "B. its", "C. it's"]),
        ("The boy asked ________ why he had failed.", ["A. he", "B. him", "C. himself"]),
        ("Some students like reading, and I am one of ________.", ["A. they", "B. them", "C. themselves"]),
        ("The technique is difficult, but people keep improving ________.", ["A. it", "B. its", "C. it's"]),
    ]:
        options(doc, q, choices)
    heading(doc, "三、短文填词（约 20 分钟）", 2)
    body(doc, "阅读短文，用括号内词的正确形式或适当的词填空。")
    passage = (
        "Nanyin is a traditional kind of music in southern Fujian. It is usually 1. ________ (play) "
        "with special instruments. Its songs are often about local life. Many young people feel like "
        "2. ________ (learn) it because it is 3. ________ good way to understand their hometown. "
        "Today, there 4. ________ many clubs for students. With the help of teachers, students can "
        "perform much 5. ________ (well) than before. They also introduce Nanyin to other 6. ________ "
        "(place). It is not easy, 7. ________ they keep practicing. By 8. ________ (work) together, "
        "they are protecting this art by 9. ________ (they). They hope it will become more 10. ________ "
        "(wide) known in the future."
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    set_font(p.add_run(passage), 11)
    heading(doc, "四、翻译固定结构（约 10 分钟）", 2)
    for q in ["继续做某事", "介意做某事", "想要做某事", "对做某事感兴趣", "花钱做某事", "通过做某事"]:
        numbered(doc, q)

    day_start(doc, 5, "综合检测与短文写作", "独立检测五天成果；完成一篇活动类短文")
    heading(doc, "一、词汇检测（约 12 分钟）", 2)
    vocab_table(doc, ["减少；降低", "目的", "正确地；恰当地", "很可能；大概", "替换；取代", "移走；去除", "研究", "打败", "表演；表现", "可用的；有空的", "勇气", "鼓励"])
    heading(doc, "二、综合单项选择（约 13 分钟）", 2)
    qs = [
        ("At that moment, I ________ a book.", ["A. read", "B. was reading", "C. am reading"]),
        ("The bridge ________ next year.", ["A. will build", "B. will be built", "C. built"]),
        ("You ________ smoke here. It is forbidden.", ["A. needn't", "B. mustn't", "C. can't be"]),
        ("She enjoys ________ English songs.", ["A. sing", "B. singing", "C. to singing"]),
        ("Tom speaks English ________ of all the students.", ["A. well", "B. better", "C. best"]),
        ("If it ________ tomorrow, we will stay at home.", ["A. rains", "B. will rain", "C. rained"]),
        ("Neither of them ________ able to come today.", ["A. are", "B. is", "C. be"]),
        ("It is ________ useful book.", ["A. a", "B. an", "C. the"]),
    ]
    for q, choices in qs:
        options(doc, q, choices)
    heading(doc, "三、句型检测（约 15 分钟）", 2)
    for q in [
        "询问父亲的职业。",
        "询问今天的天气。",
        "询问桌上有几个苹果。",
        "邀请朋友今晚和你一起看电影。",
        "请别人不要独自夜间外出。",
        "安慰一个考试失利的朋友。",
    ]:
        numbered(doc, q, 1)
    heading(doc, "四、阅读理解（约 8 分钟）", 2)
    reading = (
        "Last Saturday, our school held a Tea Culture Day. In the morning, a teacher introduced the "
        "history of tea. Then we learned how to make tea properly. At first, I thought it was simple, "
        "but I soon realized that every step required care. In the afternoon, we designed posters and "
        "shared what we had learned with younger students. I was tired but proud because the activity "
        "helped us understand traditional Chinese culture more deeply."
    )
    body(doc, reading)
    for q in [
        "When did the school hold Tea Culture Day? ____________________________________________",
        "What did the students learn after the introduction? __________________________________",
        "Why did the writer feel proud? ______________________________________________________",
        "写出文中表示“意识到”的单词：________________",
    ]:
        numbered(doc, q)
    heading(doc, "五、短文写作（约 12 分钟）", 2)
    body(doc, "假设你参加了学校的 English Culture Day。请写一篇 80—100 词的短文，介绍活动的时间、地点、两项活动以及你的感受。", bold=True)
    for _ in range(11):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(5)
        set_font(p.runs[0], 9, color=GRAY)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        set_header_footer(section)
    props = doc.core_properties
    props.title = "五天英语补缺练习册（学生版）"
    props.subject = "中考后五天轻量英语补缺练习"
    props.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
